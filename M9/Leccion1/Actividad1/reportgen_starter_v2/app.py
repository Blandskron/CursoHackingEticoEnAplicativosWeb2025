import os, shutil
from datetime import datetime
from flask import Flask, render_template, request, send_file
from pathlib import Path

# ======= Selección de motor PDF (WeasyPrint -> wkhtmltopdf -> ReportLab) =======
USE_ENGINE = os.getenv("PDF_ENGINE", "auto")  # "auto" | "weasyprint" | "wkhtmltopdf" | "reportlab"

HAS_WEASY = False
try:
    if USE_ENGINE in ("auto", "weasyprint"):
        from weasyprint import HTML, CSS  # opcional: requiere GTK/Pango (no recomendado en Windows)
        HAS_WEASY = True
except Exception:
    HAS_WEASY = False

HAS_PDFKIT = False
PDFKIT_CONFIG = None
try:
    if USE_ENGINE in ("auto", "wkhtmltopdf"):
        import pdfkit
        HAS_PDFKIT = True
        # detectar wkhtmltopdf.exe sin pedir instalación del sistema (permite bin/ local si lo agregas)
        PDFKIT_CANDIDATES = []
        env_bin = os.getenv("WKHTMLTOPDF_BIN")
        if env_bin:
            PDFKIT_CANDIDATES.append(env_bin)
        which_bin = shutil.which("wkhtmltopdf")
        if which_bin:
            PDFKIT_CANDIDATES.append(which_bin)
        # rutas típicas Windows
        PDFKIT_CANDIDATES += [
            r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe",
            r"C:\Program Files (x86)\wkhtmltopdf\bin\wkhtmltopdf.exe",
        ]
        # ruta local al proyecto (puedes poner un portable en ./bin/wkhtmltopdf.exe si quisieras)
        BASE_DIR_PROBE = Path(__file__).resolve().parent
        PDFKIT_CANDIDATES.append(str(BASE_DIR_PROBE / "bin" / "wkhtmltopdf.exe"))

        exe = next((p for p in PDFKIT_CANDIDATES if p and os.path.exists(p)), None)
        if exe:
            PDFKIT_CONFIG = pdfkit.configuration(wkhtmltopdf=exe)
        else:
            try:
                PDFKIT_CONFIG = pdfkit.configuration()  # intentará PATH
            except Exception:
                PDFKIT_CONFIG = None
except Exception:
    HAS_PDFKIT = False
    PDFKIT_CONFIG = None

# ======= ReportLab (puro Python, sin dependencias del sistema) =======
HAS_REPORTLAB = False
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image, Preformatted
    )
    HAS_REPORTLAB = True
except Exception:
    HAS_REPORTLAB = False

# ======= DOCX =======
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
BASE_DIR = Path(__file__).resolve().parent
EXPORT_DIR = BASE_DIR / "export"
EXPORT_DIR.mkdir(exist_ok=True)

DEFAULTS = {
    "titulo": "Informe de Auditoría de Seguridad 2025",
    "subtitulo": "Resultados de auditoría técnica y recomendaciones de mitigación",
    "organizacion": "Blandskron SpA",
    "autor": "Equipo de Auditoría",
    "fecha": datetime.now().strftime("%d/%m/%Y"),
    "version": "1.0",
    "clasificacion": "Confidencial",
    "resumen": (
        "Este informe presenta los hallazgos clave y las recomendaciones derivadas "
        "de la auditoría de seguridad realizada a la aplicación web institucional. El objetivo "
        "principal fue identificar vulnerabilidades críticas, evaluar su impacto y proponer "
        "medidas de mitigación. La estructura y diseño se han adaptado para ofrecer una "
        "experiencia clara, gráfica y profesional."
    ),
    "entorno": "- Superficie de ataque externa e interna.\n- API y endpoints críticos.\n- Configuraciones de servidores y bases de datos.\n- Controles de autenticación y autorización.\n- Respuesta ante errores y manejo de excepciones.",
    "tabla": [
        ["Uso de credenciales por defecto", "Acceso administrativo no autorizado, comprometiendo todo el sistema.", "Alta", "Sustituir credenciales predeterminadas, implementar MFA y rotación periódica."],
        ["Exposición de trazas internas", "Filtración de rutas y lógica interna, facilitando ataques dirigidos.", "Media", "Configurar manejo de errores genérico, almacenar trazas en sistemas seguros."],
        ["API sin autenticación", "Acceso libre a datos sensibles, incumplimiento normativo.", "Alta", "Implementar autenticación robusta (JWT/OAuth 2.0), RBAC y validaciones automáticas en CI/CD."]
    ],
    "analisis_titulo": "Análisis Detallado: API sin Autenticación",
    "analisis_descripcion": "Endpoint público expone datos personales sin requerir autenticación.",
    "analisis_herramientas": "Postman v10, cURL.",
    "analisis_metodologia": "Se realizó una solicitud GET a /api/v1/clientes sin credenciales, obteniendo datos sensibles con código HTTP 200.",
    "analisis_evidencia": "GET /api/v1/clientes HTTP/1.1\nHost: ejemplo.gov\nHTTP/1.1 200 OK\nContent-Type: application/json\n[ { \"id\": 1, \"nombre\": \"Juan Pérez\", \"email\": \"juan.perez@ejemplo.gov\" } ]",
    "analisis_recomendaciones": "- Autenticación obligatoria para todos los endpoints críticos.\n- Control de acceso granular por roles o atributos.\n- Monitorización con alertas de acceso no autorizado.",
    "conclusiones": "Las vulnerabilidades de criticidad alta representan riesgos inmediatos que deben abordarse de forma prioritaria. La implementación de controles de seguridad robustos, junto con auditorías periódicas y monitoreo continuo, es esencial para garantizar la resiliencia del entorno.",
    "recomendaciones_finales": "- Fortalecer la identidad digital: MFA resistente al phishing, revisiones de cuentas y accesos just-in-time.\n- Aumentar la visibilidad entre dominios: Integración de SIEM/XDR para correlación de eventos.\n- Defender la nube como infraestructura central: CNAPP con CDR para proteger aplicaciones nativas en la nube.\n- Priorizar vulnerabilidades: Parches regulares y gestión proactiva de riesgos.\n- Conocer al adversario: Inteligencia de amenazas y ejercicios de simulación para anticipar ataques."
}

def parse_table(raw: str):
    """Convierte texto en filas 'Hallazgo|Impacto|Criticidad|Recomendación' -> list[list[str]]"""
    rows = []
    for line in raw.strip().splitlines():
        parts = [p.strip() for p in line.split("|")]
        if len(parts) == 4:
            rows.append(parts)
    return rows

@app.route("/", methods=["GET"])
def form():
    return render_template("form.html", data=DEFAULTS)

@app.route("/preview", methods=["POST"])
def preview():
    data = dict(request.form)
    if data.get("tabla_raw"):
        parsed = parse_table(data["tabla_raw"])
        data["tabla"] = parsed if parsed else DEFAULTS["tabla"]
    else:
        data["tabla"] = DEFAULTS["tabla"]
    return render_template("report.html", data=data, preview=True)

# ======= PDF ReportLab (fallback sin deps del sistema) =======
def build_pdf_reportlab(data, out_path):
    # Imports locales para no romper si faltan módulos en otros entornos
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image, Preformatted
    )
    from reportlab.pdfgen import canvas as rl_canvas
    from pathlib import Path

    BASE_DIR_LOCAL = Path(__file__).resolve().parent

    # ---------- estilos sin colisiones ----------
    styles = getSampleStyleSheet()

    def ensure_style(name, **kwargs):
        if name in styles.byName:
            # Si existe, solo actualizamos sus atributos clave
            st = styles[name]
            for k, v in kwargs.items():
                setattr(st, k, v)
        else:
            styles.add(ParagraphStyle(name=name, **kwargs))

    ensure_style("H1x", fontSize=18, leading=22, spaceAfter=8)
    ensure_style("H2x", fontSize=14, leading=18, spaceBefore=12, spaceAfter=6)
    ensure_style("BodyJustify", fontSize=11, leading=15, alignment=TA_JUSTIFY)
    ensure_style("MetaText", fontSize=10, leading=13)
    ensure_style("CodeBlock", fontName="Courier", fontSize=9, leading=12)

    # ---------- header / footer con Página X de Y ----------
    class NumberedCanvas(rl_canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            super().showPage()

        def save(self):
            page_count = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self._draw_page_number(page_count)
                super().showPage()
            super().save()

        def _draw_page_number(self, page_count):
            self.setFont("Helvetica", 8)
            self.setFillGray(0.4)
            # pie de página
            self.drawRightString(self._pagesize[0] - 20*mm, 12*mm,
                                 f"Página {self._pageNumber} de {page_count}")

    def header_footer(canv, doc):
        canv.saveState()
        # Encabezado con logo (si existe)
        logo_path = BASE_DIR_LOCAL / "static" / "img" / "logo-horizontal.png"
        if logo_path.exists():
            try:
                canv.drawImage(str(logo_path),
                               20*mm,
                               doc.height + doc.topMargin + 5*mm,
                               width=40*mm,
                               preserveAspectRatio=True,
                               mask='auto')
            except Exception:
                pass
        canv.setFont("Helvetica", 8)
        canv.setFillGray(0.5)
        canv.drawRightString(doc.pagesize[0] - 20*mm,
                             doc.height + doc.topMargin + 12*mm,
                             "Informe generado automáticamente")
        canv.restoreState()

    # ---------- documento ----------
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=18*mm, rightMargin=18*mm, topMargin=22*mm, bottomMargin=18*mm
    )
    story = []

    # Portada
    story.append(Spacer(1, 10*mm))
    story.append(Paragraph(data.get("titulo", ""), styles["H1x"]))
    story.append(Paragraph(data.get("subtitulo", ""), styles["BodyJustify"]))
    story.append(Spacer(1, 6*mm))

    # Tabla de metadatos
    meta = [
        ["Organización", data.get("organizacion", "")],
        ["Autor", data.get("autor", "")],
        ["Fecha", data.get("fecha", "")],
        ["Versión", data.get("version", "")],
        ["Clasificación", data.get("clasificacion", "")],
    ]
    meta_tbl = Table(meta, colWidths=[40*mm, 120*mm])
    meta_tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.8, colors.black),
        ("BACKGROUND", (0, 0), (0, -1), colors.whitesmoke),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(meta_tbl)
    story.append(Spacer(1, 4*mm))
    story.append(PageBreak())

    # Secciones
    story.append(Paragraph("Resumen Ejecutivo", styles["H2x"]))
    story.append(Paragraph(data.get("resumen", "").replace("\n", "<br/>"), styles["BodyJustify"]))
    story.append(Spacer(1, 2*mm))

    story.append(Paragraph("Descripción General del Entorno", styles["H2x"]))
    story.append(Paragraph(data.get("entorno", "").replace("\n", "<br/>"), styles["BodyJustify"]))
    story.append(Spacer(1, 2*mm))

    story.append(Paragraph("Principales Vulnerabilidades Detectadas", styles["H2x"]))
    findings = [["Hallazgo", "Impacto", "Criticidad", "Recomendación"]] + data.get("tabla", [])
    widths = [40*mm, 60*mm, 20*mm, 40*mm]
    tbl = Table(findings, colWidths=widths, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.8, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e6eaf2")),
        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 2*mm))

    story.append(Paragraph(data.get("analisis_titulo", "Análisis Detallado"), styles["H2x"]))
    kv = [
        ["Descripción:", data.get("analisis_descripcion", "")],
        ["Herramientas:", data.get("analisis_herramientas", "")],
        ["Metodología:", data.get("analisis_metodologia", "")],
    ]
    kv_tbl = Table(kv, colWidths=[30*mm, 130*mm])
    kv_tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.6, colors.black),
        ("BACKGROUND", (0, 0), (0, -1), colors.whitesmoke),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(kv_tbl)
    story.append(Spacer(1, 2*mm))

    story.append(Paragraph("Evidencia", styles["H2x"]))
    story.append(Preformatted(data.get("analisis_evidencia", ""), styles["CodeBlock"]))
    story.append(Spacer(1, 2*mm))

    story.append(Paragraph("Recomendaciones", styles["H2x"]))
    story.append(Paragraph(data.get("analisis_recomendaciones", "").replace("\n", "<br/>"), styles["BodyJustify"]))
    story.append(Spacer(1, 2*mm))

    story.append(Paragraph("Conclusiones", styles["H2x"]))
    story.append(Paragraph(data.get("conclusiones", ""), styles["BodyJustify"]))
    story.append(Spacer(1, 2*mm))

    story.append(Paragraph("Recomendaciones Finales", styles["H2x"]))
    story.append(Paragraph(data.get("recomendaciones_finales", "").replace("\n", "<br/>"), styles["BodyJustify"]))

    doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer, canvasmaker=NumberedCanvas)

@app.route("/export/pdf", methods=["POST"])
def export_pdf():
    data = dict(request.form)
    if data.get("tabla_raw"):
        parsed = parse_table(data["tabla_raw"])
        data["tabla"] = parsed if parsed else DEFAULTS["tabla"]
    else:
        data["tabla"] = DEFAULTS["tabla"]

    html = render_template("report.html", data=data, preview=False)
    pdf_path = EXPORT_DIR / "Informe_Auditoria_Seguridad.pdf"

    # 1) WeasyPrint si está disponible y permitido
    if HAS_WEASY and USE_ENGINE in ("auto", "weasyprint"):
        from weasyprint import HTML, CSS  # seguro si HAS_WEASY = True
        HTML(string=html, base_url=str(BASE_DIR)).write_pdf(
            str(pdf_path),
            stylesheets=[CSS(filename=str((BASE_DIR / "static" / "css" / "styles.css").resolve()))]
        )
        return send_file(pdf_path, as_attachment=True)

    # 2) wkhtmltopdf/pdfkit si está disponible (puedes poner un portable en ./bin/)
    if HAS_PDFKIT and USE_ENGINE in ("auto", "wkhtmltopdf"):
        import pdfkit
        options = {
            "page-size": "A4",
            "margin-top": "12mm",
            "margin-right": "12mm",
            "margin-bottom": "12mm",
            "margin-left": "12mm",
            "encoding": "UTF-8",
            "print-media-type": None,
        }
        try:
            if PDFKIT_CONFIG:
                pdfkit.from_string(html, str(pdf_path), options=options, configuration=PDFKIT_CONFIG)
            else:
                pdfkit.from_string(html, str(pdf_path), options=options)
            return send_file(pdf_path, as_attachment=True)
        except Exception:
            pass  # si falla, caemos a ReportLab

    # 3) Fallback sin dependencias del sistema: ReportLab
    if HAS_REPORTLAB and USE_ENGINE in ("auto", "reportlab"):
        build_pdf_reportlab(data, pdf_path)
        return send_file(pdf_path, as_attachment=True)

    return (
        "No hay motor PDF disponible. Activa WeasyPrint (con GTK), o coloca wkhtmltopdf.exe en ./bin/, "
        "o instala reportlab y usa PDF_ENGINE=reportlab.",
        500,
    )

@app.route("/export/docx", methods=["POST"])
def export_docx():
    data = dict(request.form)
    if data.get("tabla_raw"):
        parsed = parse_table(data["tabla_raw"])
        data["tabla"] = parsed if parsed else DEFAULTS["tabla"]
    else:
        data["tabla"] = DEFAULTS["tabla"]

    doc = Document()
    # Encabezado
    section = doc.sections[0]
    header = section.header
    h_para = header.paragraphs[0]
    h_para.text = ""
    run = h_para.add_run("Blandskron | Informe de Auditoría de Seguridad")
    run.bold = True
    h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Portada
    title = doc.add_paragraph(data.get("titulo", ""))
    title.runs[0].font.size = Pt(28)
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(data.get("subtitulo", ""))
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_table(rows=5, cols=2)
    meta.style = "Table Grid"
    fields = [
        ("Organización", data.get("organizacion", "")),
        ("Autor", data.get("autor", "")),
        ("Fecha", data.get("fecha", "")),
        ("Versión", data.get("version", "")),
        ("Clasificación", data.get("clasificacion", "")),
    ]
    for i, (k, v) in enumerate(fields):
        meta.cell(i, 0).text = k
        meta.cell(i, 1).text = v

    doc.add_page_break()

    # Cuerpo
    doc.add_heading("Resumen Ejecutivo", level=1)
    p = doc.add_paragraph(data.get("resumen", ""))
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading("Descripción General del Entorno", level=1)
    for line in data.get("entorno", "").splitlines():
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_heading("Principales Vulnerabilidades Detectadas", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Hallazgo", "Impacto", "Criticidad", "Recomendación"]):
        hdr[i].text = h
    for row in data["tabla"]:
        r = table.add_row().cells
        for i, val in enumerate(row):
            r[i].text = val

    doc.add_heading(data.get("analisis_titulo", "Análisis Detallado"), level=1)
    doc.add_paragraph("Descripción: " + data.get("analisis_descripcion", ""))
    doc.add_paragraph("Herramientas: " + data.get("analisis_herramientas", ""))
    doc.add_paragraph("Metodología: " + data.get("analisis_metodologia", ""))
    doc.add_paragraph("Evidencia:\n" + data.get("analisis_evidencia", ""))
    doc.add_paragraph("Recomendaciones:\n" + data.get("analisis_recomendaciones", ""))

    doc.add_heading("Conclusiones", level=1)
    doc.add_paragraph(data.get("conclusiones", ""))

    doc.add_heading("Recomendaciones Finales", level=1)
    for line in data.get("recomendaciones_finales", "").splitlines():
        if line.strip():
            doc.add_paragraph(line.strip())

    out = EXPORT_DIR / "Informe_Auditoria_Seguridad.docx"
    doc.save(out)
    return send_file(out, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
